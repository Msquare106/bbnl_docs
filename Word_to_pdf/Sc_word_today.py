"""
pdf_to_word_layout_preserve.py
Improved scanned-PDF -> editable Word with preserved alignment/columns

Approach:
- Render page -> high-res image (PyMuPDF)
- OCR with pytesseract.image_to_data -> word-level boxes
- Cluster word x positions to form columns
- Create a Word table per page using column boundaries so alignment is preserved
- Fallback: paragraph with indent and tab stops if table not suitable

Dependencies:
    pip install pymupdf pytesseract python-docx opencv-python Pillow numpy
"""

import os
import sys
import math
import glob
import io
from collections import defaultdict, Counter

import fitz              # PyMuPDF
from PIL import Image
import pytesseract
import numpy as np
import cv2
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# ---- User configuration: adjust if needed ----
DPI = 300
MIN_COL_WIDTH_PX = 40   # minimum width of a column (px); smaller gaps become merges
WORD_CONF_THRESHOLD = 20  # ignore low confidence words
FONT_SIZE_PT = 10
# ----------------------------------------------

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"



def render_pdf_page_to_pil(pdf_path, page_number, dpi=DPI):
    doc = fitz.open(pdf_path)
    page = doc[page_number]
    mat = fitz.Matrix(dpi/72, dpi/72)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img_bytes = pix.tobytes("png")
    pil = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    doc.close()
    return pil

def ocr_words_with_boxes(pil_image):
    # pytesseract gives box coords relative to image pixels
    data = pytesseract.image_to_data(pil_image, output_type=pytesseract.Output.DICT, config='--psm 6 --oem 3')
    words = []
    n = len(data['text'])
    for i in range(n):
        txt = (data['text'][i] or "").strip()
        if not txt:
            continue
        try:
            conf = float(data['conf'][i])
        except:
            continue
        if conf < WORD_CONF_THRESHOLD:
            continue
        left, top, w, h = data['left'][i], data['top'][i], data['width'][i], data['height'][i]
        cx = left + w/2
        cy = top + h/2
        words.append({
            'text': txt,
            'conf': conf,
            'left': left,
            'top': top,
            'w': w,
            'h': h,
            'cx': cx,
            'cy': cy
        })
    return words

def detect_h_v_lines(pil_image_gray):
    img = np.array(pil_image_gray)
    if len(img.shape) == 3:
        img = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    # binary
    _, th = cv2.threshold(img, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    th_inv = 255 - th
    h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (50,1))
    v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1,50))
    hor = cv2.morphologyEx(th_inv, cv2.MORPH_OPEN, h_kernel, iterations=2)
    ver = cv2.morphologyEx(th_inv, cv2.MORPH_OPEN, v_kernel, iterations=2)
    return hor, ver

def cluster_x_positions(words, img_w):
    # Create histogram of x positions (left edges and centers)
    xs = []
    for w in words:
        xs.append(w['left'])
        xs.append(w['left'] + w['w'])
        xs.append(w['cx'])
    if not xs:
        return []
    xs = sorted(xs)
    # Agglomerative clustering by gap threshold
    clusters = []
    curr = [xs[0]]
    for x in xs[1:]:
        if x - curr[-1] <= MIN_COL_WIDTH_PX:
            curr.append(x)
        else:
            clusters.append(curr)
            curr = [x]
    if curr:
        clusters.append(curr)
    # Convert clusters to column boundaries (min left, max right)
    cols = []
    for c in clusters:
        cols.append((min(c), max(c)))
    # Normalize to page width and merge tiny columns
    merged = []
    for left, right in cols:
        if merged and left - merged[-1][1] <= MIN_COL_WIDTH_PX:
            merged[-1] = (merged[-1][0], max(merged[-1][1], right))
        else:
            merged.append((left, right))
    # Clip and ensure full page coverage
    final = []
    prev = 0
    for left, right in merged:
        l = max(0, left)
        r = min(img_w, right)
        if l - prev > MIN_COL_WIDTH_PX:
            # keep gap as empty column
            final.append((prev, l))
        final.append((l, r))
        prev = r
    if img_w - prev > MIN_COL_WIDTH_PX:
        final.append((prev, img_w))
    # Filter zero-width
    final = [(int(l), int(r)) for l, r in final if r - l > 4]
    return final

def words_to_rows_by_y(words, y_tol=10):
    # cluster words by y coordinate (line grouping)
    ys = sorted(words, key=lambda w: w['cy'])
    rows = []
    for w in ys:
        if not rows:
            rows.append([w])
        else:
            if abs(w['cy'] - np.mean([x['cy'] for x in rows[-1]])) <= y_tol:
                rows[-1].append(w)
            else:
                rows.append([w])
    # sort words inside rows by left
    for r in rows:
        r.sort(key=lambda x: x['left'])
    return rows

def build_docx_for_page(doc, pil_img, words, page_idx):
    img_w, img_h = pil_img.size

    # --- Step 1: Detect table regions via line detection ---
    gray = np.array(pil_img.convert("L"))
    _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    binary = 255 - binary

    # detect lines
    scale = 15
    horizontal = cv2.getStructuringElement(cv2.MORPH_RECT, (int(img_w/scale), 1))
    vertical   = cv2.getStructuringElement(cv2.MORPH_RECT, (1, int(img_h/scale)))

    horiz_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, horizontal, iterations=2)
    vert_lines  = cv2.morphologyEx(binary, cv2.MORPH_OPEN, vertical, iterations=2)
    table_mask = cv2.add(horiz_lines, vert_lines)

    # find table contours
    contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    table_boxes = []
    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        if w > 100 and h > 50:  # ignore small noise
            table_boxes.append((x, y, x+w, y+h))

    # --- Step 2: classify words ---
    inside, outside = [], []
    for w in words:
        cx, cy = w["cx"], w["cy"]
        found = False
        for (x1, y1, x2, y2) in table_boxes:
            if x1 <= cx <= x2 and y1 <= cy <= y2:
                inside.append(w)
                found = True
                break
        if not found:
            outside.append(w)

    # --- Step 3: process outside text first ---
    outside_rows = words_to_rows_by_y(outside, y_tol=12)
    for row in outside_rows:
        text = " ".join([w["text"] for w in row])
        left = min([w["left"] for w in row]) if row else 0
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(FONT_SIZE_PT)
        indent_in = left / img_w * 6.5
        if indent_in > 0.1:
            para.paragraph_format.left_indent = Inches(indent_in)

    # --- Step 4: process each detected table separately ---
    for (x1, y1, x2, y2) in sorted(table_boxes, key=lambda b: (b[1], b[0])):
        words_in_table = [w for w in inside if x1 <= w["cx"] <= x2 and y1 <= w["cy"] <= y2]
        if not words_in_table:
            continue
        # cluster rows/cols
        rows = words_to_rows_by_y(words_in_table, y_tol=10)
        cols = cluster_x_positions(words_in_table, img_w)

        num_rows, num_cols = len(rows), max(len(cols), 1)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        for r_idx, row in enumerate(rows):
            for w in row:
                for c_idx, (l, r) in enumerate(cols):
                    if l <= w["cx"] <= r:
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text += (" " if cell.text else "") + w["text"]
                        break

        # style
        for r in table.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(FONT_SIZE_PT)
        doc.add_paragraph()  # spacing after each table


def convert_pdf_to_docx(pdf_path, output_path=None):
    if not output_path:
        output_path = os.path.splitext(pdf_path)[0] + "_converted_layout.docx"
    doc = Document()
    # set margins smaller to better approximate page
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    pdf_doc = fitz.open(pdf_path)
    total = len(pdf_doc)
    pdf_doc.close()

    for p in range(total):
        print(f"Processing page {p+1}/{total} ...")
        pil = render_pdf_page_to_pil(pdf_path, p, dpi=DPI)
        words = ocr_words_with_boxes(pil)
        if not words:
            print("  No OCR words found on page; adding the page as image.")
            # Insert page image if OCR fails
            doc.add_paragraph()
            image_stream = io.BytesIO()
            pil.save(image_stream, format='PNG')
            image_stream.seek(0)
            doc.add_picture(image_stream, width=Inches(6.5))
            doc.add_page_break()
            continue

        build_docx_for_page(doc, pil, words, p)

        if p < total - 1:
            doc.add_page_break()

    doc.save(output_path)
    print("Saved:", output_path)
    return output_path

def main():
    if len(sys.argv) >= 2:
        target = sys.argv[1]
        if os.path.isdir(target):
            pdfs = glob.glob(os.path.join(target, "*.pdf"))
        else:
            pdfs = [target]
    else:
        pdfs = glob.glob(os.path.join(os.getcwd(), "*.pdf"))

    if not pdfs:
        print("No PDF files found in directory. Usage: python pdf_to_word_layout_preserve.py <file.pdf> or drop PDFs in current folder.")
        return

    for pdf in pdfs:
        print("Converting:", pdf)
        out = convert_pdf_to_docx(pdf)
        print("Done ->", out)

if __name__ == "__main__":
    main()
